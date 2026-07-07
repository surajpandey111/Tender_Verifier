"""
Turns (extracted_docs, verdict) into the actual deliverable a tender-desk
manager reads: one row per requirement, found/missing, extracted value,
page number to jump straight to in the source PDF, and a PASS/FAIL reason
in plain English — plus a PDF version of the same for printing/emailing.
"""

from __future__ import annotations

import os
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from docx import Document
from docx.shared import RGBColor

from reports.bidder_narrative import build_bidder_sections

STATUS_FILL = {
    "PASS": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
    "FAIL": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
    "NOT_FOUND": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
}
STATUS_COLOR_PDF = {"PASS": colors.HexColor("#2E7D32"), "FAIL": colors.HexColor("#C62828"),
                     "NOT_FOUND": colors.HexColor("#B08800")}


def generate_excel_report(tender_id: str, extracted_docs: list[dict], verdict, output_path: str):
    wb = Workbook()

    # --- Sheet 1: Eligibility summary ---
    ws1 = wb.active
    ws1.title = "Eligibility Summary"
    ws1.append([f"Tender: {tender_id}", "", "", ""])
    ws1.append([f"Overall Result: {verdict.overall_status}", "", "", ""])
    ws1.append([f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", "", "", ""])
    ws1.append([])
    header = ["Criterion", "Status", "Reason", "Evidence (page refs)"]
    ws1.append(header)
    for cell in ws1[5]:
        cell.font = Font(bold=True)

    for c in verdict.criteria:
        evidence_str = "; ".join(
            f"{e['doc_type']} p.{e['page_number']}" + (f" = {e.get('value')}" if e.get("value") is not None else "")
            for e in c.evidence
        )
        ws1.append([c.label, c.status, c.reason, evidence_str])
        row_cells = ws1[ws1.max_row]
        fill = STATUS_FILL.get(c.status)
        if fill:
            row_cells[1].fill = fill

    for i, width in enumerate([38, 12, 60, 40], start=1):
        ws1.column_dimensions[get_column_letter(i)].width = width
    for row in ws1.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # --- Sheet 2: Documents found (page-referenced checklist) ---
    ws2 = wb.create_sheet("Documents Found")
    ws2.append(["Document Type", "Source File", "Page", "Classification Method", "Confidence/Score", "Key Fields"])
    for cell in ws2[1]:
        cell.font = Font(bold=True)

    for d in extracted_docs:
        key_fields = "; ".join(f"{k}={v}" for k, v in d.get("fields", {}).items() if v)
        ws2.append([
            d.get("label", d["doc_type"]), d["source_file"], d["page_number"],
            d.get("classification_method", ""), d.get("classification_score", ""), key_fields,
        ])
    for i, width in enumerate([30, 30, 8, 18, 14, 80], start=1):
        ws2.column_dimensions[get_column_letter(i)].width = width
    for row in ws2.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(output_path)


def generate_pdf_report(tender_id: str, extracted_docs: list[dict], verdict, output_path: str):
    doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                             leftMargin=1.5 * cm, rightMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=16)
    body_style = ParagraphStyle("BodyX", parent=styles["BodyText"], fontSize=9, leading=12)

    elements = []
    elements.append(Paragraph(f"Tender Eligibility Report — {tender_id}", title_style))
    status_color = STATUS_COLOR_PDF.get(verdict.overall_status, colors.black)
    elements.append(Paragraph(
        f'Overall Result: <font color="{status_color}"><b>{verdict.overall_status}</b></font>', styles["Heading2"]
    ))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    elements.append(Spacer(1, 0.5 * cm))

    # Criteria table
    data = [["Criterion", "Status", "Reason"]]
    row_colors = []
    for c in verdict.criteria:
        data.append([Paragraph(c.label, body_style), c.status, Paragraph(c.reason, body_style)])
        row_colors.append(STATUS_COLOR_PDF.get(c.status, colors.white))

    table = Table(data, colWidths=[5 * cm, 2.2 * cm, 9.3 * cm], repeatRows=1)
    style_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]
    for i, rc in enumerate(row_colors, start=1):
        style_cmds.append(("TEXTCOLOR", (1, i), (1, i), rc))
        style_cmds.append(("FONTNAME", (1, i), (1, i), "Helvetica-Bold"))
    table.setStyle(TableStyle(style_cmds))
    elements.append(table)
    elements.append(Spacer(1, 0.8 * cm))

    # Documents found
    elements.append(Paragraph("Documents Found (page references)", styles["Heading2"]))
    doc_data = [["Document Type", "File", "Page", "Key Fields"]]
    for d in extracted_docs:
        key_fields = "; ".join(f"{k}={v}" for k, v in d.get("fields", {}).items() if v)
        doc_data.append([
            Paragraph(d.get("label", d["doc_type"]), body_style),
            Paragraph(d["source_file"], body_style),
            str(d["page_number"]),
            Paragraph(key_fields, body_style),
        ])
    doc_table = Table(doc_data, colWidths=[4 * cm, 4 * cm, 1.5 * cm, 6.9 * cm], repeatRows=1)
    doc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
    ]))
    elements.append(doc_table)

    doc.build(elements)


# ============================================================================
# Consolidated multi-bidder reports (TENDER -> many BIDDERS -> one report)
# ============================================================================
# The functions above (generate_excel_report / generate_pdf_report) produce
# ONE report per single set of documents. Real usage is different: one
# TENDER contains many BIDDER folders, and the deliverable is ONE report
# per tender listing every bidder in numbered A-G sections, matching the
# manually-written TCR format this project replaces. These functions build
# that consolidated report in PDF, XLSX, and DOCX from the same shared
# bidder_narrative.build_bidder_sections() so all three formats stay
# consistent with each other.

def generate_consolidated_pdf_report(tender_id: str, bidder_results: list[dict], output_path: str):
    """One PDF listing every bidder in numbered A-G sections, matching the sample TCR format exactly."""
    doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                             leftMargin=1.5 * cm, rightMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleX", parent=styles["Title"], fontSize=16)
    bidder_style = ParagraphStyle("BidderX", parent=styles["Heading2"], fontSize=12, spaceBefore=14)
    body_style = ParagraphStyle("BodyX", parent=styles["BodyText"], fontSize=10, leading=13, leftIndent=14)
    conclusion_style = ParagraphStyle("ConclX", parent=styles["BodyText"], fontSize=10, leading=13,
                                      leftIndent=14, textColor=colors.HexColor("#B00020"))

    elements = [Paragraph(f"Technical Compliance Report — {tender_id}", title_style),
                Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | {len(bidder_results)} bidder(s)", styles["Normal"]),
                Spacer(1, 0.4 * cm)]

    for i, br in enumerate(bidder_results, start=1):
        s = build_bidder_sections(br["bidder_name"], br["extracted_docs"], br["verdict"])
        elements.append(Paragraph(f"{i}) {s['bidder_name']}", bidder_style))

        elements.append(Paragraph("A) Legal Status:", body_style))
        for line in s["legal_status"]:
            elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{line}", body_style))
        elements.append(Paragraph(f"B) Declaration (Banned/Suspended/Blacklisted):&nbsp;&nbsp;{s['declaration']}", body_style))
        elements.append(Paragraph(f"C) Technical Specification:&nbsp;&nbsp;{s['technical_spec']}", body_style))
        elements.append(Paragraph(f"D) MSE / Udyam:&nbsp;&nbsp;{s['mse']}", body_style))
        elements.append(Paragraph(f"E) EMD:&nbsp;&nbsp;{s['emd']}", body_style))
        elements.append(Paragraph(f"F) Turnover:&nbsp;&nbsp;{s['turnover']}", body_style))
        elements.append(Paragraph("G) Experience:", body_style))
        for line in s["experience"]:
            elements.append(Paragraph(f"&nbsp;&nbsp;&nbsp;{line}", body_style))

        elements.append(Paragraph("CONCLUSION:", body_style))
        for line in s["conclusion"].split("\n"):
            elements.append(Paragraph(line.replace("- ", "&nbsp;&nbsp;&nbsp;- "), conclusion_style))
        elements.append(Spacer(1, 0.3 * cm))

    doc.build(elements)


def generate_consolidated_xlsx_report(tender_id: str, bidder_results: list[dict], output_path: str):
    """Summary sheet (one row per bidder) + Details sheet (full A-G text per bidder)."""
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Summary"
    header = ["#", "Bidder", "GSTIN", "PAN", "UDIN", "Avg Turnover", "Experience Certs Found", "Result", "Conclusion"]
    ws1.append(header)
    for cell in ws1[1]:
        cell.font = Font(bold=True)

    for i, br in enumerate(bidder_results, start=1):
        s = build_bidder_sections(br["bidder_name"], br["extracted_docs"], br["verdict"])
        raw = s["raw"]
        ws1.append([i, s["bidder_name"], raw["gstin"], raw["pan_number"], raw["udin"], raw["avg_turnover"],
                    raw["experience_count"], s["overall_status"], s["conclusion"]])
        fill = STATUS_FILL.get(s["overall_status"])
        if fill:
            ws1.cell(row=ws1.max_row, column=8).fill = fill

    for i, width in enumerate([4, 30, 20, 14, 22, 16, 12, 12, 60], start=1):
        ws1.column_dimensions[get_column_letter(i)].width = width
    for row in ws1.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    ws2 = wb.create_sheet("Details")
    ws2.append(["#", "Bidder", "A) Legal Status", "B) Declaration", "C) Tech Spec", "D) MSE/Udyam",
                "E) EMD", "F) Turnover", "G) Experience", "Conclusion"])
    for cell in ws2[1]:
        cell.font = Font(bold=True)
    for i, br in enumerate(bidder_results, start=1):
        s = build_bidder_sections(br["bidder_name"], br["extracted_docs"], br["verdict"])
        ws2.append([i, s["bidder_name"], "\n".join(s["legal_status"]), s["declaration"], s["technical_spec"],
                    s["mse"], s["emd"], s["turnover"], "\n".join(s["experience"]), s["conclusion"]])
    for i, width in enumerate([4, 30, 35, 20, 20, 25, 20, 30, 45, 45], start=1):
        ws2.column_dimensions[get_column_letter(i)].width = width
    for row in ws2.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(output_path)


def generate_consolidated_docx_report(tender_id: str, bidder_results: list[dict], output_path: str):
    """Word version of the same numbered A-G report, for whoever needs to edit/annotate it further in Word."""
    doc = Document()
    doc.add_heading(f"Technical Compliance Report — {tender_id}", level=1)
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | {len(bidder_results)} bidder(s)")
    meta.runs[0].italic = True

    for i, br in enumerate(bidder_results, start=1):
        s = build_bidder_sections(br["bidder_name"], br["extracted_docs"], br["verdict"])

        doc.add_heading(f"{i}) {s['bidder_name']}", level=2)

        doc.add_paragraph().add_run("A) Legal Status:").bold = True
        for line in s["legal_status"]:
            doc.add_paragraph(line, style="List Bullet")

        p = doc.add_paragraph()
        p.add_run("B) Declaration (Banned/Suspended/Blacklisted): ").bold = True
        p.add_run(s["declaration"])

        p = doc.add_paragraph()
        p.add_run("C) Technical Specification: ").bold = True
        p.add_run(s["technical_spec"])

        p = doc.add_paragraph()
        p.add_run("D) MSE / Udyam: ").bold = True
        p.add_run(s["mse"])

        p = doc.add_paragraph()
        p.add_run("E) EMD: ").bold = True
        p.add_run(s["emd"])

        p = doc.add_paragraph()
        p.add_run("F) Turnover: ").bold = True
        p.add_run(s["turnover"])

        doc.add_paragraph().add_run("G) Experience:").bold = True
        for line in s["experience"]:
            doc.add_paragraph(line, style="List Bullet")

        doc.add_paragraph().add_run("CONCLUSION:").bold = True
        for line in s["conclusion"].split("\n"):
            concl_p = doc.add_paragraph(line)
            for run in concl_p.runs:
                run.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)

        if i < len(bidder_results):
            doc.add_page_break()

    doc.save(output_path)


def generate_consolidated_reports(tender_id: str, bidder_results: list[dict], output_dir: str) -> dict:
    """Convenience wrapper: writes report.pdf, report.xlsx, report.docx into output_dir. Returns their paths."""
    paths = {
        "pdf": os.path.join(output_dir, "report.pdf"),
        "xlsx": os.path.join(output_dir, "report.xlsx"),
        "docx": os.path.join(output_dir, "report.docx"),
    }
    generate_consolidated_pdf_report(tender_id, bidder_results, paths["pdf"])
    generate_consolidated_xlsx_report(tender_id, bidder_results, paths["xlsx"])
    generate_consolidated_docx_report(tender_id, bidder_results, paths["docx"])
    return paths
